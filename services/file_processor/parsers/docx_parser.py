"""
DOCX text extraction using python-docx.
Requirements: 5.2, 5.4
"""
from __future__ import annotations
import io
import re
import docx
from docx.oxml.ns import qn


def _normalize(text: str) -> str:
    text = re.sub(r"[^\x20-\x7E\n]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file preserving heading hierarchy."""
    document = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if "Heading" in style_name:
            parts.append(f"\n\n{text}")
        else:
            parts.append(text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return _normalize("\n".join(parts))
